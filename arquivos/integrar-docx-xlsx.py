'''
9. Gerar relatórios em Word a partir de Excel
● Crie um programa que:
○ Leia uma planilha Excel com dados (ex.: "Nome do Cliente", "Valor da Compra").
○ Crie um arquivo Word chamado relatorio.docx.
○ Adicione os dados do Excel no Word, formatados em forma de tabela.

#socorro Deus
'''

import openpyxl
from docx import Document
from openpyxl import load_workbook

wb = load_workbook('vendas.xlsx')
ws = wb.active

doc = Document()
doc.add_heading('Relatório de Vendas', level=1)

tabela = doc.add_table(rows=1, cols=ws.max_column)
tabela.style = 'Table Grid'


for coluna, cell in enumerate(ws[1], start=0):
    tabela.cell(0, coluna).text = str(cell.value)


for row in ws.iter_rows(min_row=2, values_only=True):  
    linha = tabela.add_row().cells
    for coluna, value in enumerate(row):
        linha[coluna].text = str(value)


doc.save('relatorio.docx')
print('Relatório gerado...')

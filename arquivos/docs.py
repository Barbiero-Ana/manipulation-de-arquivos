'''
1. Criar um arquivo Word
● Crie um programa que:
○ Gera um arquivo Word chamado documento.docx.
○ Adiciona um título em negrito e sublinhado.
○ Adiciona dois parágrafos de texto.
○ Salva o arquivo.

* Lembrar de tentar implementar um sistema de menu nesse código, provavel que vai ser em def AAAAAAAAAAAAAAAAA
'''

from docx import Document
from docx.enum.text import WD_UNDERLINE #biblioteca que deixa o titulo estilizado (sublinhado ou em negrito, essas coisas ai)



doc = Document()
titulo = doc.add_heading('Curso de Python na CyberEdux', level = 0)
#o level= 0 deixa o texto grande? euacho

for run in titulo.runs:
    run.bold = True
    run.underline = WD_UNDERLINE.SINGLE

# titulo.bold = True
# titulo.underline = True

doc.add_paragraph('Como funciona o curso de python? Bem, vamos la.')
doc.add_paragraph('O curso funciona em turmas de manhã, tarde e noite explicando o funcionamento do python e suas funcionalidades dentro da área de programação')
doc.save('documento.docx')

# 2 - Abra o arquivo word já existente e leia seu conteúdo e o imprima

doc = Document('documento.docx')
for paragraph in doc.paragraphs:
    print(paragraph.text)

# 3 - Faça com que leia o arquivo word, conte o número total de palavras no arquivo e imprima o resultado

doc = Document('documento.docx')
cont_palavras = 0
for paragraph in doc.paragraphs:
    palavras = paragraph.text.split()
    cont_palavras += len(palavras)
    print(f'O total de palavras no arquivo é: {cont_palavras}')  

# 4 - Gere um novo arquivo doc sobre tabela e insira uma tabela com 3 colunas e 5 linhas - preencha as células com valores sla linha x e coluna y e salve o arquivo :D

doc = Document()
table = doc.add_table( rows= 5, cols= 3)

for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        cell.add_paragraph(f'Linha {i + 1}, Coluna {j + 1}')
doc.save('tabela.docx')

# 5 - Leia um arquivo já existente e substitua uma palavras seleciona (em especifico né) e troque todas as suas ocorrências dentro do arquivo por uma outra palavras, tipo 'python' troque por 'BOMDIA' e claro, salve o arquivo....


doc = Document('documento.docx')
while True:
    print('\nDeseja substituir alguma palavra dentro do arquivo docx\n1 - Sim\n2 - Não')
    decisao = int(input('- '))

    if decisao == 1:
        palavra = input('Qual palavra deseja substituir: ')
        palavra_new = input('Digite a palavra que irá substituila: ')

        doc = Document('documento.docx')
        for i in doc.paragraphs:
            if palavra in i.text:
                i.text = i.text.replace(palavra,palavra_new)

        doc.save('novo_documento.docx')

    elif decisao == 2:
        print('Ok, continuando o sistema.')
        break
    else:
        print('Opção inválida')
    


